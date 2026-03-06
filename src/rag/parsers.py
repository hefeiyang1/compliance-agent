"""
Document parsers for ComplianceAgent RAG system.

This module provides parsers for various document formats including PDF, Word documents,
and Markdown files. Each parser extracts text content and metadata from documents.
"""

import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from enum import Enum

import pypdf
import pdfplumber
from docx import Document as DocxDocument

from loguru import logger


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "md"
    TEXT = "txt"


@dataclass
class ParsedDocument:
    """
    Represents a parsed document with its content and metadata.

    Attributes:
        content: The extracted text content
        metadata: Additional metadata extracted from the document
        page_count: Number of pages (for PDFs) or sections
        total_tokens: Estimated token count
        document_type: Type of document that was parsed
    """
    content: str
    metadata: Dict[str, any]
    page_count: int
    total_tokens: int
    document_type: DocumentType

    def calculate_hash(self) -> str:
        """Calculate SHA-256 hash of the content for deduplication."""
        return hashlib.sha256(self.content.encode('utf-8')).hexdigest()

    def estimate_tokens(self) -> int:
        """
        Estimate token count using a simple heuristic.

        Uses the rule of thumb that 1 token ≈ 4 characters for English text.
        For Chinese text, 1 token ≈ 1.5 characters.
        """
        # Count Chinese characters
        chinese_chars = sum(1 for c in self.content if '\u4e00' <= c <= '\u9fff')
        # Count other characters
        other_chars = len(self.content) - chinese_chars

        # Estimate tokens
        estimated_tokens = (chinese_chars / 1.5) + (other_chars / 4)
        return int(estimated_tokens)


class BaseParser:
    """Base class for document parsers."""

    def __init__(self, file_path: str):
        """
        Initialize the parser.

        Args:
            file_path: Path to the document file
        """
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

    def parse(self) -> ParsedDocument:
        """
        Parse the document and extract content.

        Returns:
            ParsedDocument object with content and metadata

        Raises:
            NotImplementedError: Must be implemented by subclasses
        """
        raise NotImplementedError("Subclasses must implement parse()")

    def _get_file_metadata(self) -> Dict[str, any]:
        """Extract basic file metadata."""
        stat = self.file_path.stat()
        return {
            "filename": self.file_path.name,
            "file_size": stat.st_size,
            "extension": self.file_path.suffix,
        }


class PDFParser(BaseParser):
    """
    Parser for PDF documents using pdfplumber and pypdf.

    This parser uses pdfplumber as the primary parser for better text extraction,
    with pypdf as a fallback.
    """

    def __init__(self, file_path: str, use_pdfplumber: bool = True):
        """
        Initialize the PDF parser.

        Args:
            file_path: Path to the PDF file
            use_pdfplumber: Whether to use pdfplumber (True) or pypdf (False)
        """
        super().__init__(file_path)
        if self.file_path.suffix.lower() not in ['.pdf']:
            raise ValueError(f"Not a PDF file: {file_path}")
        self.use_pdfplumber = use_pdfplumber

    def parse(self) -> ParsedDocument:
        """
        Parse the PDF document.

        Returns:
            ParsedDocument with extracted text and metadata
        """
        if self.use_pdfplumber:
            return self._parse_with_pdfplumber()
        else:
            return self._parse_with_pypdf()

    def _parse_with_pdfplumber(self) -> ParsedDocument:
        """Parse PDF using pdfplumber for better text extraction."""
        content_parts = []
        metadata = self._get_file_metadata()
        page_texts = []

        try:
            with pdfplumber.open(self.file_path) as pdf:
                metadata.update({
                    "page_count": len(pdf.pages),
                    "parser": "pdfplumber",
                })

                # Extract text from each page
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            page_texts.append(text)
                            content_parts.append(f"--- Page {page_num} ---\n{text}\n")
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {e}")
                        content_parts.append(f"--- Page {page_num} ---\n[Failed to extract text]\n")

                # Try to extract metadata from PDF info
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    pdf_metadata = {}
                    for key, value in pdf.metadata.items():
                        if value and isinstance(value, str):
                            pdf_metadata[key] = value
                    if pdf_metadata:
                        metadata["pdf_metadata"] = pdf_metadata

        except Exception as e:
            logger.error(f"Failed to parse PDF with pdfplumber: {e}")
            # Fallback to pypdf
            return self._parse_with_pypdf()

        content = "\n".join(content_parts)

        # Create parsed document
        parsed_doc = ParsedDocument(
            content=content,
            metadata=metadata,
            page_count=len(page_texts),
            total_tokens=0,  # Will be calculated
            document_type=DocumentType.PDF
        )
        parsed_doc.total_tokens = parsed_doc.estimate_tokens()

        logger.info(f"Parsed PDF with {len(page_texts)} pages, ~{parsed_doc.total_tokens} tokens")
        return parsed_doc

    def _parse_with_pypdf(self) -> ParsedDocument:
        """Parse PDF using pypdf as fallback."""
        content_parts = []
        metadata = self._get_file_metadata()
        page_texts = []

        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)

                metadata.update({
                    "page_count": len(pdf_reader.pages),
                    "parser": "pypdf",
                })

                # Extract metadata if available
                if pdf_reader.metadata:
                    pdf_metadata = {}
                    for key, value in pdf_reader.metadata.items():
                        if value and isinstance(value, str):
                            # Clean up metadata keys (remove leading '/')
                            clean_key = key.lstrip('/')
                            pdf_metadata[clean_key] = value
                    if pdf_metadata:
                        metadata["pdf_metadata"] = pdf_metadata

                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            page_texts.append(text)
                            content_parts.append(f"--- Page {page_num} ---\n{text}\n")
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {e}")
                        content_parts.append(f"--- Page {page_num} ---\n[Failed to extract text]\n")

        except Exception as e:
            logger.error(f"Failed to parse PDF with pypdf: {e}")
            raise

        content = "\n".join(content_parts)

        # Create parsed document
        parsed_doc = ParsedDocument(
            content=content,
            metadata=metadata,
            page_count=len(page_texts),
            total_tokens=0,  # Will be calculated
            document_type=DocumentType.PDF
        )
        parsed_doc.total_tokens = parsed_doc.estimate_tokens()

        logger.info(f"Parsed PDF with {len(page_texts)} pages, ~{parsed_doc.total_tokens} tokens")
        return parsed_doc


class DocxParser(BaseParser):
    """
    Parser for Microsoft Word documents (.docx).

    Extracts text content and preserves some structure like paragraphs.
    """

    def __init__(self, file_path: str):
        """
        Initialize the DOCX parser.

        Args:
            file_path: Path to the Word document
        """
        super().__init__(file_path)
        if self.file_path.suffix.lower() not in ['.docx']:
            raise ValueError(f"Not a Word document: {file_path}")

    def parse(self) -> ParsedDocument:
        """
        Parse the Word document.

        Returns:
            ParsedDocument with extracted text and metadata
        """
        try:
            doc = DocxDocument(str(self.file_path))

            # Extract basic metadata
            metadata = self._get_file_metadata()
            metadata.update({
                "parser": "python-docx",
                "paragraph_count": len(doc.paragraphs),
            })

            # Extract core properties if available
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.created:
                metadata["created"] = doc.core_properties.created.isoformat()
            if doc.core_properties.modified:
                metadata["modified"] = doc.core_properties.modified.isoformat()

            # Extract text from paragraphs
            content_parts = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    content_parts.append(text)

            # Extract text from tables
            if doc.tables:
                table_count = 0
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        content_parts.append(table_text)
                        table_count += 1
                metadata["table_count"] = table_count

            content = "\n\n".join(content_parts)

            # Create parsed document
            parsed_doc = ParsedDocument(
                content=content,
                metadata=metadata,
                page_count=metadata["paragraph_count"],  # Use paragraphs as page analog
                total_tokens=0,
                document_type=DocumentType.DOCX
            )
            parsed_doc.total_tokens = parsed_doc.estimate_tokens()

            logger.info(f"Parsed DOCX with {metadata['paragraph_count']} paragraphs, ~{parsed_doc.total_tokens} tokens")
            return parsed_doc

        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise

    def _extract_table_text(self, table) -> Optional[str]:
        """
        Extract text from a table as a formatted string.

        Args:
            table: python-docx Table object

        Returns:
            Formatted table text or None if empty
        """
        if not table.rows:
            return None

        rows_text = []
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells_text))

        table_text = "\n".join(rows_text)
        return f"[Table]\n{table_text}\n[/Table]"


class MarkdownParser(BaseParser):
    """
    Parser for Markdown files.

    Preserves Markdown structure and extracts metadata from frontmatter if present.
    """

    def __init__(self, file_path: str, preserve_structure: bool = True):
        """
        Initialize the Markdown parser.

        Args:
            file_path: Path to the Markdown file
            preserve_structure: Whether to preserve Markdown formatting
        """
        super().__init__(file_path)
        if self.file_path.suffix.lower() not in ['.md', '.markdown']:
            raise ValueError(f"Not a Markdown file: {file_path}")
        self.preserve_structure = preserve_structure

    def parse(self) -> ParsedDocument:
        """
        Parse the Markdown file.

        Returns:
            ParsedDocument with extracted text and metadata
        """
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            metadata = self._get_file_metadata()
            metadata.update({
                "parser": "markdown",
                "preserve_structure": self.preserve_structure,
            })

            # Extract frontmatter if present (YAML metadata between ---)
            lines = content.split('\n')
            if lines and lines[0].strip() == '---':
                # Find the end of frontmatter
                end_idx = 1
                for i, line in enumerate(lines[1:], 1):
                    if line.strip() == '---':
                        end_idx = i
                        break

                # Parse YAML frontmatter
                frontmatter_lines = lines[1:end_idx]
                if frontmatter_lines:
                    frontmatter = self._parse_frontmatter(frontmatter_lines)
                    if frontmatter:
                        metadata["frontmatter"] = frontmatter
                        # Extract common fields
                        if "title" in frontmatter:
                            metadata["title"] = frontmatter["title"]
                        if "author" in frontmatter:
                            metadata["author"] = frontmatter["author"]
                        if "date" in frontmatter:
                            metadata["date"] = frontmatter["date"]

            # Count sections (headers)
            section_count = sum(1 for line in lines if line.startswith('#'))
            metadata["section_count"] = section_count

            # If not preserving structure, remove some Markdown formatting
            if not self.preserve_structure:
                # Remove header markers but keep text
                content = self._remove_markdown_formatting(content)

            # Create parsed document
            parsed_doc = ParsedDocument(
                content=content,
                metadata=metadata,
                page_count=section_count,  # Use sections as page analog
                total_tokens=0,
                document_type=DocumentType.MARKDOWN
            )
            parsed_doc.total_tokens = parsed_doc.estimate_tokens()

            logger.info(f"Parsed Markdown with {section_count} sections, ~{parsed_doc.total_tokens} tokens")
            return parsed_doc

        except Exception as e:
            logger.error(f"Failed to parse Markdown: {e}")
            raise

    def _parse_frontmatter(self, lines: List[str]) -> Dict[str, any]:
        """
        Parse YAML frontmatter into a dictionary.

        Simple parser that handles key: value pairs.

        Args:
            lines: Frontmatter lines (without --- markers)

        Returns:
            Dictionary of metadata
        """
        frontmatter = {}
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                key = key.strip()
                value = value.strip()
                # Remove quotes if present
                if value.startswith('"') and value.endswith('"'):
                    value = value[1:-1]
                elif value.startswith("'") and value.endswith("'"):
                    value = value[1:-1]
                frontmatter[key] = value

        return frontmatter

    def _remove_markdown_formatting(self, content: str) -> str:
        """
        Remove some Markdown formatting while preserving content.

        Args:
            content: Raw Markdown content

        Returns:
            Content with reduced formatting
        """
        lines = content.split('\n')
        cleaned_lines = []

        for line in lines:
            # Remove header markers but keep the text
            if line.startswith('#'):
                line = line.lstrip('#').strip()

            # Remove bold/italic markers
            line = line.replace('**', '').replace('__', '')
            line = line.replace('*', '').replace('_', '')

            # Remove link syntax but keep text and URL
            while '[text](url)' in line:
                import re
                line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)

            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)


class TextParser(BaseParser):
    """
    Parser for plain text files.

    Simple parser that reads text files and extracts basic metadata.
    """

    def __init__(self, file_path: str):
        """
        Initialize the text parser.

        Args:
            file_path: Path to the text file
        """
        super().__init__(file_path)
        if self.file_path.suffix.lower() not in ['.txt', '.text']:
            raise ValueError(f"Not a text file: {file_path}")

    def parse(self) -> ParsedDocument:
        """
        Parse the text file.

        Returns:
            ParsedDocument with extracted text and metadata
        """
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            metadata = self._get_file_metadata()
            metadata.update({
                "parser": "text",
                "line_count": len(content.split('\n')),
            })

            # Create parsed document
            parsed_doc = ParsedDocument(
                content=content,
                metadata=metadata,
                page_count=metadata["line_count"],  # Use lines as page analog
                total_tokens=0,
                document_type=DocumentType.TEXT
            )
            parsed_doc.total_tokens = parsed_doc.estimate_tokens()

            logger.info(f"Parsed text file with {metadata['line_count']} lines, ~{parsed_doc.total_tokens} tokens")
            return parsed_doc

        except Exception as e:
            logger.error(f"Failed to parse text file: {e}")
            raise


def create_parser(file_path: str, **kwargs) -> BaseParser:
    """
    Factory function to create the appropriate parser based on file extension.

    Args:
        file_path: Path to the document file
        **kwargs: Additional arguments to pass to the parser

    Returns:
        Appropriate parser instance

    Raises:
        ValueError: If file type is not supported
    """
    path = Path(file_path)
    extension = path.suffix.lower()

    parser_map = {
        '.pdf': PDFParser,
        '.docx': DocxParser,
        '.md': MarkdownParser,
        '.markdown': MarkdownParser,
        '.txt': TextParser,
        '.text': TextParser,
    }

    parser_class = parser_map.get(extension)
    if parser_class is None:
        raise ValueError(f"Unsupported file type: {extension}. Supported types: {list(parser_map.keys())}")

    return parser_class(file_path, **kwargs)


def parse_document(file_path: str, **kwargs) -> ParsedDocument:
    """
    Convenience function to parse a document.

    Args:
        file_path: Path to the document file
        **kwargs: Additional arguments to pass to the parser

    Returns:
        ParsedDocument with content and metadata

    Example:
        >>> doc = parse_document("compliance_policy.pdf")
        >>> print(f"Content: {doc.content[:100]}")
        >>> print(f"Tokens: {doc.total_tokens}")
    """
    parser = create_parser(file_path, **kwargs)
    return parser.parse()
